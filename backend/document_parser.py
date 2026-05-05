"""
document_parser.py — Text extraction pipeline for uploaded user documents.

Supports:
  PDF      : embedded-text extraction via pdfplumber; OCR fallback via
             pytesseract + pdf2image for scanned documents.
  Image    : OCR via pytesseract (JPG, PNG, HEIC)
  DOCX     : python-docx
  TXT      : plain read
  CSV      : pandas — each row → labeled field line
  XLSX/XLS : pandas — all sheets extracted and labeled by sheet name

System requirements (must be installed at container level):
  apt-get install -y tesseract-ocr poppler-utils
  See /root/.emergent/on-restart.sh for the persistence hook.
"""

import logging
import subprocess
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Magic bytes: XLSX / ZIP files always start with PK\x03\x04
_XLSX_MAGIC = b'PK\x03\x04'

# MIME types so generic that we can't trust them — fall through to ext-sniff
_GENERIC_MIMES = frozenset({
    'application/octet-stream',
    'binary/octet-stream',
    'application/x-binary',
    '',
})

# ── Supported MIME types ─────────────────────────────────────────────────────
SUPPORTED_MIME_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    "image/png": "png",
    "image/heic": "heic",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    # ── Spreadsheets (Prompt: spreadsheet support) ──────────────────────────
    "text/csv": "csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
}

MAX_FILE_BYTES        = 10 * 1024 * 1024   # 10 MB
MIN_PDF_TEXT_LEN      = 100                # below this we assume scanned PDF
_MAX_ROWS_PER_SHEET   = 5_000              # row cap for spreadsheet parsing


# ── Startup capability check ────────────────────────────────────────────────────
def check_parse_capabilities() -> dict:
    """
    Run once at server startup. Returns a dict describing the availability
    of each parse dependency.  'ok' is False if anything critical is missing.
    """
    result: dict = {"ok": True, "details": {}}

    # Check tesseract binary
    try:
        proc = subprocess.run(
            ["tesseract", "--version"],
            capture_output=True, text=True, timeout=5,
        )
        ver_line = (proc.stdout or proc.stderr or "").strip().split("\n")[0]
        result["details"]["tesseract"] = ver_line if proc.returncode == 0 else "error"
        if proc.returncode != 0:
            result["ok"] = False
    except FileNotFoundError:
        result["details"]["tesseract"] = "NOT FOUND"
        result["ok"] = False

    # Check pdftoppm binary (poppler)
    try:
        proc = subprocess.run(
            ["pdftoppm", "-v"],
            capture_output=True, text=True, timeout=5,
        )
        ver_line = (proc.stderr or proc.stdout or "").strip().split("\n")[0]
        result["details"]["pdftoppm"] = ver_line or "ok"
    except FileNotFoundError:
        result["details"]["pdftoppm"] = "NOT FOUND"
        result["ok"] = False

    # Check Python libraries
    lib_map = {
        "pypdf":       "pypdf",
        "pdfplumber":  "pdfplumber",
        "pdf2image":   "pdf2image",
        "pytesseract": "pytesseract",
        "python-docx": "docx",
        "Pillow":      "PIL",
        # Spreadsheet parsing — pandas + openpyxl (for .xlsx)
        "pandas":      "pandas",
        "openpyxl":    "openpyxl",
    }
    for display_name, import_name in lib_map.items():
        try:
            __import__(import_name)
            result["details"][display_name] = "ok"
        except ImportError:
            result["details"][display_name] = "NOT INSTALLED"
            # pandas/openpyxl missing is non-fatal for existing formats
            if import_name not in ("pandas", "openpyxl"):
                result["ok"] = False

    return result


# ── Parse helpers ──────────────────────────────────────────────────────────────────

def _extract_pdf_text(file_path: str) -> tuple[str, int]:
    """Extract text from a PDF. Falls back to OCR if embedded text is sparse."""
    import pdfplumber

    text_parts: list[str] = []
    page_count = 0
    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text() or ""
                text_parts.append(t)
    except Exception as exc:
        logger.warning("pdfplumber failed on %s: %s", file_path, exc)

    full_text = "\n".join(text_parts).strip()

    if len(full_text) < MIN_PDF_TEXT_LEN:
        logger.info(
            "PDF has only %d chars of embedded text — falling back to OCR",
            len(full_text),
        )
        return _ocr_pdf(file_path)

    return full_text, page_count


def _ocr_pdf(file_path: str) -> tuple[str, int]:
    """Rasterise PDF pages and run OCR on each."""
    from pdf2image import convert_from_path
    import pytesseract

    try:
        images = convert_from_path(file_path, dpi=200)
    except Exception as exc:
        raise RuntimeError(f"pdf2image conversion failed: {exc}") from exc

    texts = [pytesseract.image_to_string(img) or "" for img in images]
    return "\n\n".join(texts).strip(), len(images)


def _extract_image_text(file_path: str) -> tuple[str, int]:
    """Run OCR on a single image file."""
    import pytesseract
    from PIL import Image

    img = Image.open(file_path)
    # Convert HEIC or unusual modes to RGB so tesseract is happy
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    text = pytesseract.image_to_string(img) or ""
    return text.strip(), 1


def _extract_docx_text(file_path: str) -> tuple[str, int]:
    """Extract plain text from a .docx file."""
    import docx

    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs)
    page_count = max(1, len(doc.sections))
    return text, page_count


def _extract_txt(file_path: str) -> tuple[str, int]:
    """Read a plain-text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        text = fh.read()
    return text, 1


# ── Spreadsheet parsers ──────────────────────────────────────────────────────

def _row_to_line(row: "pd.Series", cols: list) -> str:  # type: ignore[name-defined]
    """Convert a DataFrame row to a labeled-field string, skipping empty cells."""
    parts = [f"{col}: {row[col]}" for col in cols if str(row[col]).strip()]
    return " | ".join(parts)


def _extract_csv_text(file_path: str) -> tuple[str, int]:
    """
    Parse a CSV file and convert each row to a labeled field line.
    e.g.  Day: Monday | Exercise: Squat | Sets: 5 | Reps: 3 | Weight: 405
    Caps at _MAX_ROWS_PER_SHEET rows; logs and truncates if exceeded.
    """
    import pandas as pd

    df = pd.read_csv(file_path, dtype=str, keep_default_na=False)
    if len(df) > _MAX_ROWS_PER_SHEET:
        logger.warning("CSV has %d rows — truncating to %d", len(df), _MAX_ROWS_PER_SHEET)
        df = df.head(_MAX_ROWS_PER_SHEET)

    cols  = list(df.columns)
    lines = [_row_to_line(row, cols) for _, row in df.iterrows()]
    text  = "\n".join(line for line in lines if line)
    return text, 1


def _extract_xlsx_text(file_path: str) -> tuple[str, int]:
    """
    Parse an XLSX or XLS file. All sheets are extracted and labeled with
    their sheet name as a section header.
    Caps each sheet at _MAX_ROWS_PER_SHEET rows.
    """
    import pandas as pd

    xl       = pd.ExcelFile(file_path)
    sections = []

    for sheet_name in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str, keep_default_na=False)
        if df.empty:
            continue
        if len(df) > _MAX_ROWS_PER_SHEET:
            logger.warning(
                "Sheet '%s' has %d rows — truncating to %d",
                sheet_name, len(df), _MAX_ROWS_PER_SHEET,
            )
            df = df.head(_MAX_ROWS_PER_SHEET)

        cols  = list(df.columns)
        lines = [_row_to_line(row, cols) for _, row in df.iterrows()]
        body  = "\n".join(line for line in lines if line)
        if body:
            sections.append(f"--- Sheet: {sheet_name} ---\n{body}")

    text = "\n\n".join(sections)
    return text, len(xl.sheet_names)


# ── Public entry point ──────────────────────────────────────────────────────────────────


def _is_xlsx_magic(file_path: str) -> bool:
    """Return True if the file starts with the XLSX/ZIP magic bytes PK\\x03\\x04."""
    try:
        with open(file_path, 'rb') as fh:
            return fh.read(4) == _XLSX_MAGIC
    except Exception:
        return False


# File-extension → parser key for the fallback lookup
_EXT_DISPATCH: dict[str, str] = {
    '.pdf':  'pdf',
    '.docx': 'docx',
    '.doc':  'docx',
    '.txt':  'txt',
    '.csv':  'csv',
    '.xlsx': 'xlsx',
    '.xls':  'xlsx',   # sniff on dispatch below
    '.png':  'img',
    '.jpg':  'img',
    '.jpeg': 'img',
    '.heic': 'img',
}


def _dispatch(key: str, file_path: str) -> tuple[str, int]:
    """Route a resolved parser key to the correct extraction function."""
    if key == 'pdf':   return _extract_pdf_text(file_path)
    if key == 'img':   return _extract_image_text(file_path)
    if key == 'docx':  return _extract_docx_text(file_path)
    if key == 'txt':   return _extract_txt(file_path)
    if key == 'csv':   return _extract_csv_text(file_path)
    if key == 'xlsx':  return _extract_xlsx_text(file_path)
    raise ValueError(f"Unknown parser key: {key!r}")


def parse_document(file_path: str, content_type: str) -> tuple[str, int]:
    """
    Extract text from the given file.  Returns (text, page_count).

    Dispatch strategy
    -----------------
    1. MIME-type dispatch for well-known content types.
    2. Special case: application/vnd.ms-excel  →  sniff magic bytes.
       - ZIP/XLSX header (PK\\x03\\x04)  →  xlsx parser
       - anything else                 →  csv parser
       (Windows often labels CSVs with this MIME type.)
    3. If content_type is generic/missing/octet-stream, or if MIME lookup
       produced no match, fall back to the file's lowercased extension.
    4. Raises ValueError only if BOTH MIME and extension lookups fail.
    """
    ct = (content_type or '').lower().strip()

    # ── Special case: sniff the ambiguous vnd.ms-excel MIME ──────────────────
    if ct == 'application/vnd.ms-excel':
        if _is_xlsx_magic(file_path):
            logger.info("parse_document: vnd.ms-excel + XLSX magic → xlsx parser")
            return _extract_xlsx_text(file_path)
        logger.info("parse_document: vnd.ms-excel without XLSX magic → csv parser")
        return _extract_csv_text(file_path)

    # ── MIME-type dispatch (non-generic types only) ───────────────────────────
    if ct and ct not in _GENERIC_MIMES:
        mime_map: dict[str, str] = {
            'application/pdf':                                                      'pdf',
            'image/jpeg':                                                           'img',
            'image/jpg':                                                            'img',
            'image/png':                                                            'img',
            'image/heic':                                                           'img',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain':                                                           'txt',
            'text/csv':                                                             'csv',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':   'xlsx',
        }
        key = mime_map.get(ct)
        if key:
            return _dispatch(key, file_path)
        # Unrecognised but non-generic MIME — fall through to extension lookup
        logger.info(
            "parse_document: unrecognised MIME %r → trying extension fallback",
            content_type,
        )

    # ── Extension fallback ────────────────────────────────────────────────────
    ext = Path(file_path).suffix.lower()
    logger.info(
        "parse_document: content_type=%r (generic or unmatched) → ext fallback, ext=%r",
        content_type, ext,
    )
    key = _EXT_DISPATCH.get(ext)
    if key:
        return _dispatch(key, file_path)

    raise ValueError(
        f"Unsupported content type {content_type!r} (extension: {ext!r}). "
        f"Supported types: pdf, docx, txt, csv, xlsx/xls, png, jpg/jpeg, heic."
    )
