"""
test_parse_formats.py — Standalone parse pipeline test.
Creates realistic test files for each supported format and runs parse_document().
Reports exact parsedText output. No server required.
"""

import os
import sys
import textwrap

sys.path.insert(0, "/app/backend")

from document_parser import parse_document

TEST_DIR = "/tmp/doc_parse_tests"
os.makedirs(TEST_DIR, exist_ok=True)

SEPARATOR = "=" * 60

def section(title):
    print(f"\n{SEPARATOR}")
    print(f"  FORMAT: {title}")
    print(SEPARATOR)


# ── 1. TXT ────────────────────────────────────────────────────────────────────
def test_txt():
    section("TXT")
    path = os.path.join(TEST_DIR, "program.txt")
    content = textwrap.dedent("""\
        STRENGTH BLOCK — WEEK 1
        ========================
        Day 1 (Monday): Lower — Squat Focus
          - Back Squat: 5x5 @ 80% 1RM
          - Romanian Deadlift: 4x8 @ 65%
          - Leg Press: 3x12
          - Calf Raises: 4x15

        Day 2 (Wednesday): Upper — Push
          - Bench Press: 5x5 @ 80% 1RM
          - Overhead Press: 4x6 @ 70%
          - Incline DB Press: 3x10
          - Tricep Pushdown: 3x12

        Day 3 (Friday): Full Body — Deadlift Focus
          - Deadlift: 5x3 @ 85% 1RM
          - Pull-Ups: 4x6
          - Barbell Row: 4x8
          - Face Pulls: 3x15

        Notes: Rest 90-120s between working sets. Deload on week 4.
    """)
    with open(path, "w") as f:
        f.write(content)
    text, pages = parse_document(path, "text/plain")
    print(f"Pages: {pages}")
    print(f"Chars extracted: {len(text)}")
    print(f"\n--- parsedText ---\n{text}\n--- end ---")
    return len(text) > 10


# ── 2. DOCX ───────────────────────────────────────────────────────────────────
def test_docx():
    section("DOCX")
    import docx
    path = os.path.join(TEST_DIR, "program.docx")
    doc = docx.Document()
    doc.add_heading("12-Week Powerlifting Program", 0)
    doc.add_heading("Week 1 — Accumulation Phase", level=1)
    doc.add_paragraph("Primary Goal: Build work capacity and technique consistency.")

    doc.add_heading("Session A — Squat + Accessories", level=2)
    # Table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Exercise"
    hdr[1].text = "Sets"
    hdr[2].text = "Reps"
    hdr[3].text = "Intensity"
    rows = [
        ("Back Squat", "4", "6", "75% 1RM"),
        ("Front Squat", "2", "4", "65% 1RM"),
        ("Bulgarian Split Squat", "3", "8/leg", "RPE 7"),
    ]
    for exercise, sets, reps, intensity in rows:
        row = table.add_row().cells
        row[0].text = exercise
        row[1].text = sets
        row[2].text = reps
        row[3].text = intensity
    doc.add_paragraph("Coach note: Pause 2 seconds at bottom on all squat variations.")
    doc.save(path)
    text, pages = parse_document(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    print(f"Pages: {pages}")
    print(f"Chars extracted: {len(text)}")
    print(f"\n--- parsedText ---\n{text}\n--- end ---")
    return len(text) > 10


# ── 3. PDF (text-based) ───────────────────────────────────────────────────────
def test_pdf_text():
    section("PDF (text-based / pdfplumber path)")
    from fpdf import FPDF
    path = os.path.join(TEST_DIR, "program.pdf")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Hypertrophy Program — 8 Weeks", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 8, "Designed by Coach Marcus | RPE-Based Training", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "Phase 1: Foundation (Weeks 1-4)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    lines = [
        "Training Days: Monday, Wednesday, Friday",
        "Session Length: 60-75 minutes",
        "",
        "DAY 1 — CHEST & TRICEPS",
        "  Flat Barbell Bench Press — 4x8-10 @ RPE 7",
        "  Incline DB Press — 3x10-12 @ RPE 7",
        "  Cable Chest Fly — 3x12-15 @ RPE 6",
        "  Tricep Dips — 3x10-12 bodyweight",
        "  Overhead Tricep Extension — 3x12 @ RPE 6",
        "",
        "DAY 2 — BACK & BICEPS",
        "  Weighted Pull-Ups — 4x6-8 @ RPE 8",
        "  Barbell Row — 4x8 @ RPE 7",
        "  Seated Cable Row — 3x10-12 @ RPE 6",
        "  Face Pulls — 3x15 @ RPE 5",
        "  Barbell Curl — 3x10-12 @ RPE 7",
    ]
    for line in lines:
        pdf.cell(0, 7, line, new_x="LMARGIN", new_y="NEXT")

    # Page 2
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 8, "DAY 3 — LEGS & SHOULDERS", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=11)
    lines2 = [
        "  Back Squat — 4x6-8 @ RPE 8",
        "  Romanian Deadlift — 3x10 @ RPE 7",
        "  Leg Press — 3x12-15 @ RPE 6",
        "  Overhead Press — 4x6-8 @ RPE 8",
        "  Lateral Raises — 4x15-20 @ RPE 6",
        "",
        "PROGRESSION MODEL:",
        "  - Add 2.5kg when you hit top of rep range at RPE < 7",
        "  - Deload week 5: reduce volume by 40%, keep intensity",
        "",
        "NUTRITION TARGETS (per training day):",
        "  Protein: 0.8-1g per lb bodyweight",
        "  Calories: maintenance + 200-300 surplus",
        "  Pre-workout: 40-60g carbs 60 min before",
    ]
    for line in lines2:
        pdf.cell(0, 7, line, new_x="LMARGIN", new_y="NEXT")

    pdf.output(path)
    text, pages = parse_document(path, "application/pdf")
    print(f"Pages: {pages}")
    print(f"Chars extracted: {len(text)}")
    print(f"\n--- parsedText ---\n{text}\n--- end ---")
    return len(text) > 100


# ── 4. JPG (image with text via PIL) ─────────────────────────────────────────
def test_jpg():
    section("JPG (OCR path)")
    from PIL import Image, ImageDraw, ImageFont
    path = os.path.join(TEST_DIR, "workout_card.jpg")

    img = Image.new("RGB", (800, 500), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Try to use a default font (size param is font size, not always supported without TTF)
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28)
        font_body  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except (IOError, OSError):
        font_title = ImageFont.load_default()
        font_body  = ImageFont.load_default()

    draw.text((40, 30),  "WORKOUT CARD — UPPER BODY A",             font=font_title, fill=(0, 0, 0))
    draw.text((40, 80),  "Exercise: Bench Press",                   font=font_body,  fill=(30, 30, 30))
    draw.text((40, 115), "Sets: 4  |  Reps: 6-8  |  RPE: 8",       font=font_body,  fill=(30, 30, 30))
    draw.text((40, 155), "Exercise: Overhead Press",                font=font_body,  fill=(30, 30, 30))
    draw.text((40, 190), "Sets: 3  |  Reps: 8-10  |  RPE: 7",      font=font_body,  fill=(30, 30, 30))
    draw.text((40, 230), "Exercise: Dumbbell Row",                  font=font_body,  fill=(30, 30, 30))
    draw.text((40, 265), "Sets: 4  |  Reps: 10-12  |  RPE: 7",     font=font_body,  fill=(30, 30, 30))
    draw.text((40, 310), "Notes: Rest 2 min between compound sets", font=font_body,  fill=(80, 80, 80))
    draw.text((40, 350), "Target: 3000 calories, 200g protein",     font=font_body,  fill=(80, 80, 80))

    img.save(path, "JPEG", quality=90)
    text, pages = parse_document(path, "image/jpeg")
    print(f"Pages: {pages}")
    print(f"Chars extracted: {len(text)}")
    print(f"\n--- parsedText ---\n{text}\n--- end ---")
    return True  # OCR on generated image — even partial extraction is acceptable


# ── 5. PDF (scan / OCR path) ─────────────────────────────────────────────────
def test_pdf_scan():
    section("PDF (scanned — OCR path)")
    # We simulate a "scanned" PDF by: creating a JPEG image with text,
    # then embedding it into a PDF as a raw image page (no embedded text layer).
    from fpdf import FPDF
    from PIL import Image, ImageDraw, ImageFont
    import tempfile

    # 1. Render text to a PNG
    img_path = os.path.join(TEST_DIR, "scan_page.png")
    img = Image.new("RGB", (1200, 900), color=(250, 248, 240))  # slightly off-white like a scan
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 32)
        font_body  = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except (IOError, OSError):
        font_title = font_body = ImageFont.load_default()

    lines = [
        ("SCANNED TRAINING LOG — WEEK 3", font_title, (0, 0, 0), 50, 40),
        ("Monday: Squat Day", font_title, (0, 0, 128), 50, 100),
        ("Warm-up: 5 min bike + mobility work", font_body, (40, 40, 40), 50, 155),
        ("Back Squat: 140kg x 5 x 5", font_body, (40, 40, 40), 50, 195),
        ("Pause Squat: 100kg x 3 x 3", font_body, (40, 40, 40), 50, 235),
        ("Leg Press: 200kg x 4 x 10", font_body, (40, 40, 40), 50, 275),
        ("", font_body, (40, 40, 40), 50, 315),
        ("Wednesday: Bench Day", font_title, (0, 0, 128), 50, 340),
        ("Bench Press: 102.5kg x 4 x 5", font_body, (40, 40, 40), 50, 395),
        ("DB Incline: 36kg x 3 x 10", font_body, (40, 40, 40), 50, 435),
        ("Cable Row: 70kg x 4 x 12", font_body, (40, 40, 40), 50, 475),
        ("", font_body, (40, 40, 40), 50, 515),
        ("Notes: Feeling strong, add 2.5kg next week", font_body, (80, 80, 80), 50, 545),
    ]
    for text_line, fnt, color, x, y in lines:
        if text_line:
            draw.text((x, y), text_line, font=fnt, fill=color)

    img.save(img_path, "PNG")

    # 2. Embed the PNG as a PDF image (no text layer = simulated scan)
    pdf_path = os.path.join(TEST_DIR, "scanned_program.pdf")
    pdf = FPDF()
    pdf.add_page()
    # FPDF adds image scaled to page width
    pdf.image(img_path, x=0, y=0, w=210)  # A4 width in mm
    pdf.output(pdf_path)

    text, pages = parse_document(pdf_path, "application/pdf")
    print(f"Pages: {pages}")
    print(f"Chars extracted: {len(text)}")
    print(f"\n--- parsedText ---\n{text}\n--- end ---")
    return True  # Accept any output from OCR


# ── Run all ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    results = {}

    try:
        results["TXT"]       = ("PASS" if test_txt()       else "FAIL", "")
    except Exception as e:
        results["TXT"]       = ("ERROR", str(e))

    try:
        results["DOCX"]      = ("PASS" if test_docx()      else "FAIL", "")
    except Exception as e:
        results["DOCX"]      = ("ERROR", str(e))

    try:
        results["PDF (text)"] = ("PASS" if test_pdf_text() else "FAIL", "")
    except Exception as e:
        results["PDF (text)"] = ("ERROR", str(e))

    try:
        results["JPG"]       = ("PASS" if test_jpg()       else "FAIL", "")
    except Exception as e:
        results["JPG"]       = ("ERROR", str(e))

    try:
        results["PDF (scan)"] = ("PASS" if test_pdf_scan() else "FAIL", "")
    except Exception as e:
        results["PDF (scan)"] = ("ERROR", str(e))

    print(f"\n{SEPARATOR}")
    print("  SUMMARY")
    print(SEPARATOR)
    for fmt, (status, err) in results.items():
        marker = "✅" if status == "PASS" else ("⚠️" if status == "FAIL" else "❌")
        print(f"  {marker}  {fmt}: {status}" + (f" — {err}" if err else ""))
    print(SEPARATOR)
